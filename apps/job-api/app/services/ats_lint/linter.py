"""Deterministic ATS format linter over rendered `.docx` bytes.

Rules (Greenhouse-floor):
- valid_docx (error)          — opens as a valid zipped OOXML package
- no_tables (error)           — `doc.tables` empty
- no_inline_images (error)    — no `word/media/*` entries in the zip
- no_text_frames (error)      — no `txbx` in document XML
- no_shapes (error)           — no `w:drawing` / `w:pict` in document XML
- experience_heading (error)  — at least one Heading 1 named "Experience"
- standard_headings (warning) — every Heading 1 in {Summary, Experience, Skills, Education}
- page_count (warning/error)  — ~paragraph count heuristic: > 80 warn, > 120 error

The renderer module already avoids these failure modes by construction.
The linter is the safety net — it catches the case where the renderer
is extended and accidentally reintroduces one.
"""

from __future__ import annotations

import io
import zipfile

from docx import Document

from app.models.ats_lint import LintResult, LintViolation

_ALLOWED_HEADING_1_NAMES = {"Summary", "Experience", "Skills", "Education"}
_PAGE_COUNT_WARN_AT = 80
_PAGE_COUNT_ERROR_AT = 120


def _parse_or_none(data: bytes) -> object | None:
    try:
        return Document(io.BytesIO(data))
    except Exception:
        return None


def _raw_xml_or_empty(doc: object) -> str:
    try:
        return str(doc.element.xml)  # type: ignore[attr-defined]
    except Exception:
        return ""


def _zip_entries(data: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            return zf.namelist()
    except zipfile.BadZipFile:
        return []


def _heading_1_texts(doc: object) -> list[str]:
    out: list[str] = []
    for para in doc.paragraphs:  # type: ignore[attr-defined]
        style = getattr(para.style, "name", "") or ""
        if style == "Heading 1":
            out.append(para.text)
    return out


def _non_empty_paragraph_count(doc: object) -> int:
    paragraphs = doc.paragraphs  # type: ignore[attr-defined]
    return len([p for p in paragraphs if p.text.strip()])


def lint_docx(data: bytes) -> LintResult:
    violations: list[LintViolation] = []

    entries = _zip_entries(data)
    doc = _parse_or_none(data)
    if doc is None or not entries:
        violations.append(
            LintViolation(
                code="valid_docx",
                message="Bytes are not a valid .docx archive.",
                severity="error",
            )
        )
        return LintResult(ok=False, violations=violations)

    # no_inline_images
    if any(name.startswith("word/media/") for name in entries):
        violations.append(
            LintViolation(
                code="no_inline_images",
                message=(
                    "Document contains inline images (word/media/ entries). "
                    "ATS parsers routinely drop image content."
                ),
                severity="error",
            )
        )

    # no_tables
    if len(doc.tables) > 0:  # type: ignore[attr-defined]
        violations.append(
            LintViolation(
                code="no_tables",
                message=(
                    "Document contains tables. Most ATS parsers read tables "
                    "inconsistently; use single-column paragraph content."
                ),
                severity="error",
            )
        )

    xml = _raw_xml_or_empty(doc)
    # no_text_frames
    if "txbx" in xml:
        violations.append(
            LintViolation(
                code="no_text_frames",
                message=(
                    "Document contains text frames (txbx). Text inside frames "
                    "often does not survive ATS parsing."
                ),
                severity="error",
            )
        )
    # no_shapes
    if "w:drawing" in xml or "w:pict" in xml:
        violations.append(
            LintViolation(
                code="no_shapes",
                message=(
                    "Document contains drawings or picture shapes. These are "
                    "unreadable to ATS parsers."
                ),
                severity="error",
            )
        )

    headings = _heading_1_texts(doc)
    # experience_heading
    if "Experience" not in headings:
        violations.append(
            LintViolation(
                code="experience_heading",
                message=(
                    'Missing required Heading 1 "Experience". ATS parsers key '
                    "off standard section names."
                ),
                severity="error",
            )
        )

    # standard_headings
    non_standard = [h for h in headings if h not in _ALLOWED_HEADING_1_NAMES]
    if non_standard:
        violations.append(
            LintViolation(
                code="standard_headings",
                message=(
                    f"Non-standard Heading 1(s): {sorted(set(non_standard))}. "
                    f"Expected a subset of {sorted(_ALLOWED_HEADING_1_NAMES)}."
                ),
                severity="warning",
            )
        )

    # page_count
    paragraph_count = _non_empty_paragraph_count(doc)
    if paragraph_count > _PAGE_COUNT_ERROR_AT:
        violations.append(
            LintViolation(
                code="page_count",
                message=(
                    f"Document has {paragraph_count} non-empty paragraphs; "
                    f"likely exceeds 3 pages. Tighten content."
                ),
                severity="error",
            )
        )
    elif paragraph_count > _PAGE_COUNT_WARN_AT:
        violations.append(
            LintViolation(
                code="page_count",
                message=(
                    f"Document has {paragraph_count} non-empty paragraphs; "
                    f"likely exceeds 2 pages."
                ),
                severity="warning",
            )
        )

    ok = not any(v.severity == "error" for v in violations)
    return LintResult(ok=ok, violations=violations)
